from docx import Document
import sys
doc = Document(r'd:\Cydra Social\CYDRA Social – Pricing Comparison.docx')
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)
for t in doc.tables:
    print('--- TABLE ---')
    for row in t.rows:
        print(' | '.join([cell.text for cell in row.cells]))
