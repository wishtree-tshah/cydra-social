import sys
print("Starting...", file=sys.stderr)

try:
    from docx import Document
    print("docx imported successfully", file=sys.stderr)
    
    doc = Document(r'd:\Cydra Social\Cydra Social - User Stories.docx')
    print(f"Document loaded. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}", file=sys.stderr)
    
    # Extract all paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(text)
    
    # Extract tables if any
    for table in doc.tables:
        print("\n--- TABLE ---")
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            print(" | ".join(row_text))
            
except Exception as e:
    print(f"Error: {e}", file=sys.stderr)
    import traceback
    traceback.print_exc()
